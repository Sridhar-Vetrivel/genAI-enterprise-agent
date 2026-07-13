"""Builds RFP_Solution_Proposal_Psiog_Kendra.docx for the GenAI Enterprise Copilot RFP.

Follows the structure of docs/RFP_Proposal_Sample_Model_ParticipantsV2.docx (Part B —
Participant Proposal), populated with content from docs/solution-proposal.md and
docs/use-case.md. Uses the AgentField approach only.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


OUT = Path(__file__).resolve().parents[1] / "docs" / "RFP_Solution_Proposal_Psiog_Kendra.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def build():
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ===== Title =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("RFP Solution Proposal — Psiog Kendra")
    r.bold = True
    r.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("GenAI-Powered Enterprise Copilot for Psiog")
    sr.italic = True
    sr.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Response to RFP S4-I-24  ·  Impact pSiddhi pSiddhi-2026-01  ·  "
        "Semester 4 Capstone  ·  Budget Ceiling ₹2,500 (Fixed)"
    ).italic = True

    doc.add_paragraph()

    # ===== Part B header =====
    add_heading(doc, "PART B — Proposal Submitted by Participant (the Vendor)", level=1)
    add_para(
        doc,
        "Use Case Proposal — Response to RFP: S4-I-24 — GenAI-Powered Enterprise Copilot (Psiog Kendra)",
        italic=True,
    )

    # ===== 1. Topic ID & Title =====
    add_heading(doc, "1. Topic ID & Title", level=2)
    add_table(
        doc,
        ["Field", "Detail"],
        [
            ["Topic ID", "S4-I-24"],
            ["Topic Title", "GenAI-Powered Enterprise Copilot"],
            ["Project Name", "Psiog Kendra"],
        ],
    )

    # ===== 2. Participant Details =====
    add_heading(doc, "2. Participant Details", level=2)
    add_table(
        doc,
        ["Field", "Detail"],
        [
            ["Full Name", "Sridhar Vetrivel"],
            ["Email", "tac@psiog.com"],
            ["Program", "Impact pSiddhi — pSiddhi-2026-01"],
            ["Semester", "Semester 4 — Capstone"],
            ["Submitted To", "L&D Team (acting as Customer)"],
            ["Proposal Date", "End of Week 2 (2026-05-19)"],
        ],
    )

    # ===== 3. Semester & Category =====
    add_heading(doc, "3. Semester & Category", level=2)
    add_para(
        doc,
        "Semester 4 — Capstone (Multi-Agent GenAI Solution, QA Mandatory, AI Core). "
        "This project — Psiog Kendra — is an end-to-end multi-agent enterprise copilot, "
        "delivered as a working POC by Week 17 within the ₹2,500 budget ceiling.",
    )

    # ===== 4. Problem Understanding =====
    add_heading(doc, "4. Problem Understanding", level=2)
    add_para(
        doc,
        "Psiog's engineering, data, and operations teams work across a fragmented technology estate — "
        "Databricks for data pipelines, GitHub/Azure DevOps for deployments, a CRM system for customer "
        "data, and scattered internal documentation across runbooks and architecture docs. These systems "
        "are disconnected — each requires separate logins, separate navigation, and separate mental models.",
    )
    add_para(doc, "This fragmentation creates three compounding problems:", bold=True)
    add_bullet(
        doc,
        "Onboarding bottleneck — new joiners cannot independently answer basic operational questions "
        "for months because the knowledge of which system to look in, and how to query it, is tribal — "
        "it lives in senior engineers' heads.",
    )
    add_bullet(
        doc,
        "Senior engineer drain — senior engineers become human routers, constantly interrupted to "
        "answer questions that are fundamentally lookups across these systems.",
    )
    add_bullet(
        doc,
        "Decision lag — operational decisions are slowed because synthesizing a complete picture "
        "(did the pipeline succeed, is the runbook up to date, did it affect any customer) requires "
        "manually visiting multiple tools and stitching results together.",
    )
    add_para(
        doc,
        "The problem is not a lack of data — Psiog has the data. The problem is that there is no "
        "unified, intelligent interface that can access all of it in response to a plain-English question.",
    )

    add_para(doc, "Why existing approaches fall short:", bold=True)
    add_table(
        doc,
        ["Approach", "Why It Fails"],
        [
            ["Asking a senior engineer", "Doesn't scale; blocks the expert; knowledge is lost when they leave"],
            ["Separate dashboards per system", "Still requires the user to know which dashboard, how to query it, and how to correlate results manually"],
            ["A single-LLM chatbot", "Answers from training data, not live system data — produces confident but unverifiable responses"],
            ["Basic keyword-based search", "Cannot understand intent, cannot cross-reference systems, cannot synthesize a unified answer"],
        ],
    )

    add_para(doc, "Success condition:", bold=True)
    add_para(
        doc,
        "A new joiner on Day 1 can ask Psiog Kendra the same question a senior engineer would know "
        "by memory, and get a correct, cited answer — without needing to know which system to look in, "
        "what credentials to use, or how to navigate any of those platforms.",
        italic=True,
    )

    # ===== 5. Proposed Solution =====
    add_heading(doc, "5. Proposed Solution", level=2)
    add_para(
        doc,
        "We propose building Psiog Kendra — a multi-agent enterprise copilot with a natural-language "
        "interface that orchestrates across four system domains: Data Platform (Databricks), DevOps "
        "(GitHub / Azure DevOps), CRM, and Internal Documentation.",
    )
    add_para(
        doc,
        "The system uses a Coordinator Agent that receives every user query, determines intent, and "
        "routes it to one or more Specialist Agents. Each Specialist Agent is responsible for exactly "
        "one domain — it queries the live system, retrieves relevant data, and returns a cited response. "
        "The Coordinator synthesizes multi-agent results into a single, grounded reply.",
    )

    add_para(doc, "Architecture (conceptual):", bold=True)
    arch = (
        "User (any employee)\n"
        "        │  Natural-language query\n"
        "        ▼\n"
        "┌─────────────────────────────────────┐\n"
        "│         Coordinator Agent           │\n"
        "│  1. Classifies query intent (LLM)   │\n"
        "│  2. Identifies target domain(s)     │\n"
        "│  3. Dispatches to specialist(s)     │\n"
        "│  4. Synthesizes a grounded reply    │\n"
        "└─────────────────────────────────────┘\n"
        "        │\n"
        "   ┌────┴───────────────┬──────────────┬──────────────┐\n"
        "   ▼                    ▼              ▼              ▼\n"
        "Data Agent          DevOps Agent   CRM Agent     Docs Agent\n"
        "Databricks REST   GitHub/Azure   CRM REST API  AgentField\n"
        "(Jobs, SQL WH)    DevOps REST    (HubSpot/mock) Vector RAG\n"
        "   │                    │              │              │\n"
        "   └────────────────────┴──────────────┴──────────────┘\n"
        "                         │\n"
        "              Grounded, cited response\n"
        "                         ▼\n"
        "              Employee sees answer with citations"
    )
    p = doc.add_paragraph()
    r = p.add_run(arch)
    r.font.name = "Consolas"
    r.font.size = Pt(9)

    add_para(doc, "Agent definitions:", bold=True)
    add_table(
        doc,
        ["Agent", "Role", "Data Source"],
        [
            ["Coordinator", "Query classifier and orchestrator — routes to specialists, synthesizes final response", "LLM (via OpenRouter)"],
            ["Data Platform Agent", "Databricks specialist — pipeline status, job history, analytics", "Databricks REST API (Jobs, SQL Warehouses)"],
            ["DevOps Agent", "CI/CD specialist — build status, deployment history, quality gates", "GitHub Actions / Azure DevOps REST API"],
            ["CRM Agent", "Customer-data specialist — contact records, deal status, account history", "CRM REST API (HubSpot free / mock)"],
            ["Docs Agent", "Internal-knowledge specialist — semantic search over runbooks and arch docs", "AgentField vector memory (RAG)"],
        ],
    )

    add_para(doc, "Why AgentField as the orchestration framework:", bold=True)
    add_para(
        doc,
        "The RFP suggests LangChain + Microsoft Agent Framework for orchestration, Azure AI Search "
        "for semantic search, and Azure Functions for serverless execution. We propose replacing this "
        "stack with AgentField — an open-source (Apache 2.0) AI backend framework — because a single "
        "framework handles routing, tracing, and agent-to-agent calls natively, includes built-in "
        "vector memory (no separate search service), and runs on a free-tier Azure VM via Docker. "
        "Every agent call is automatically traced and visualised as a DAG, providing built-in "
        "observability that would otherwise require additional tooling.",
    )
    add_table(
        doc,
        ["Suggested Tool", "AgentField Replacement", "Reason"],
        [
            ["LangChain + MS Agent Framework", "AgentField SDK orchestration", "Single framework — no glue code between two separate frameworks"],
            ["Azure AI Search", "AgentField vector memory fabric", "Built-in set_vector + similarity_search — no external service, no cost"],
            ["Azure Functions", "AgentField control plane (Docker on Azure VM)", "Simpler deployment, free-tier VM sufficient for POC scale"],
        ],
    )

    add_para(doc, "State and memory design:", bold=True)
    add_table(
        doc,
        ["Scope", "What Is Stored", "Example"],
        [
            ["session", "Conversation history for multi-turn queries", "Previous questions in a session"],
            ["global", "Pre-indexed documentation embeddings", "Runbook embeddings, doc chunks"],
            ["run", "Intermediate results within a single query", "Partial responses from each specialist"],
        ],
    )

    # ===== 6. Tools, Subscriptions & Cost Breakdown =====
    add_heading(doc, "6. Tools, Subscriptions & Cost Breakdown", level=2)
    add_para(
        doc,
        "Total estimated spend: ₹2,500 — all infrastructure components run on free tiers; the "
        "₹2,500 ceiling is fully allocated to LLM inference, which is the sole paid item.",
        italic=True,
    )
    add_table(
        doc,
        ["Tool / Service", "Purpose", "Tier", "Cost/Sem", "Justification"],
        [
            ["AgentField (self-hosted)", "Multi-agent orchestration + vector memory", "Apache 2.0 OSS", "₹0", "Single framework replaces LangChain + MS Agent FW + Azure AI Search"],
            ["OpenRouter (LLM gateway)", "LLM inference for routing, response gen, judge", "Pay-per-use", "₹2,500", "Budget ceiling fully allocated — the sole paid component"],
            ["Ollama + Llama 4 Scout", "Local LLM fallback during development", "Free (local)", "₹0", "Runs on dev machine, zero API cost; full fallback if budget approaches limit"],
            ["Databricks Free Edition", "Data platform — pipelines and jobs", "Free tier", "₹0", "Free tier covers development and demo workloads"],
            ["GitHub API + Actions", "DevOps data source + CI for tests", "Free tier", "₹0", "Public API; Actions: 2,000 min/month free"],
            ["HubSpot Free CRM / mock", "CRM data source", "Free tier", "₹0", "HubSpot free tier or realistic mock data for POC"],
            ["Azure B1s VM", "Hosts AgentField control plane (Docker)", "Free tier (12-month)", "₹0", "750 hrs/month B1s — sufficient for POC"],
            ["AgentField vector memory", "Semantic search over internal docs (RAG)", "Built-in", "₹0", "Included in AgentField — replaces Azure AI Search"],
            ["GitHub Copilot Free", "Dev-time acceleration", "Free tier", "₹0", "Free tier for development"],
            ["Pytest + GitHub Actions", "QA test suite + CI", "OSS / free", "₹0", "Industry-standard, zero cost"],
            ["TOTAL", "", "", "₹2,500", "100% of ceiling — LLM inference is the only cost item"],
        ],
    )
    add_para(
        doc,
        "Budget risk management: every non-LLM tool is either open-source self-hosted or free-tier "
        "cloud. If OpenRouter spend approaches the ceiling during development, Ollama + Llama 4 Scout "
        "provides a complete local fallback at ₹0, with no code change required (model swap via env var).",
    )

    # ===== 7. Timeline & Effort =====
    add_heading(doc, "7. Timeline & Effort", level=2)
    add_para(doc, "Development spans Weeks 4–17. Proposal submitted end of Week 2.", italic=True)

    add_para(doc, "Phase 1: Weeks 4–10 — Foundation, Specialist Agents, Mid-Term Checkpoint", bold=True)
    add_table(
        doc,
        ["Week", "Task", "Deliverable", "QA Activity"],
        [
            ["Wk 4", "AgentField control plane up via Docker Compose; Azure VM provisioned; OpenRouter connected", "Infra ready", "Smoke tests for control plane + LLM call"],
            ["Wk 5", "Coordinator Agent — intent classification + routing logic; Pydantic schemas (RoutingDecision, AgentResponse)", "Coordinator routing to named domains", "Unit tests for routing classifier"],
            ["Wk 6", "Data Platform Agent — Databricks REST API (Jobs, SQL Warehouses)", "Job-run status / history queries working", "Unit + integration tests for Databricks skills"],
            ["Wk 7", "Docs Agent + RAG index — chunk, embed, store in AgentField vector memory", "Similarity search returning relevant results", "Coverage tests for indexing + retrieval"],
            ["Wk 8", "DevOps Agent — GitHub Actions / Azure DevOps API", "Build status + deployment + quality gate queries", "Integration tests per skill"],
            ["Wk 9", "CRM Agent — HubSpot / mock CRM API", "Contact, deal, account queries", "Integration tests"],
            ["Wk 10", "Week 10 Checkpoint — 4 agents live; routing across all domains; 6 of 12 test queries passing with citations", "Mid-term docs on Moodle", "Unit tests ≥50% coverage"],
        ],
    )

    add_para(doc, "Phase 2: Weeks 11–17 — Cross-Domain Synthesis, AI-QA, Deployment, Final Demo", bold=True)
    add_table(
        doc,
        ["Week", "Task", "Deliverable", "QA Activity"],
        [
            ["Wk 11", "Cross-domain queries — parallel agent dispatch + session memory", "Multi-domain routing working", "Cross-domain E2E tests"],
            ["Wk 12", "Full test suite — all 12 queries; routing accuracy measured", "Initial hallucination rate documented", "Routing accuracy ≥ target"],
            ["Wk 13", "Judge Agent (AI QA) — grounding verification reasoner", "Automated hallucination detection per run", "Judge agent validation tests"],
            ["Wk 14", "Azure deployment — full system on B1s VM end-to-end", "Cloud deployment live", "Deployment smoke + regression"],
            ["Wk 15", "QA hardening — ≥80% coverage; edge cases; response quality validated", "Coverage targets met", "Full regression suite green"],
            ["Wk 16", "Demo prep — demo script + evidence pack (routing logs, hallucination rates, test results)", "Demo-ready system + artefacts", "Final regression run"],
            ["Wk 17", "Week 17 Final Demo — full Psiog Kendra live; all 12 queries passing; hallucination rate documented", "Final submission on Moodle", "Full QA report"],
        ],
    )
    add_para(doc, "Total effort: ~150 hours across 14 weeks (~11 hrs/week). QA is embedded each week.", italic=True)

    # ===== 8. Tech Stack =====
    add_heading(doc, "8. Tech Stack", level=2)
    add_table(
        doc,
        ["Layer", "Technologies", "Cost"],
        [
            ["Language", "Python 3.11", "Free"],
            ["Agent Orchestration", "AgentField (Apache 2.0) — @app.reasoner, @app.skill, app.call", "Free"],
            ["LLM Inference (primary)", "OpenRouter (model swappable via env var; e.g., google/gemini-2.5-flash)", "₹2,500"],
            ["LLM Inference (fallback)", "Ollama + Llama 4 Scout (local)", "Free"],
            ["Vector Memory / Semantic Search", "AgentField built-in (set_vector + similarity_search)", "Free"],
            ["Data Platform API", "Databricks REST API (Jobs, SQL Warehouses)", "Free tier"],
            ["DevOps APIs", "GitHub Actions API / Azure DevOps REST API", "Free tier"],
            ["CRM API", "HubSpot Free CRM or mock REST service", "Free"],
            ["Hosting", "Azure B1s VM (12-month free) running Docker Compose", "Free tier"],
            ["Schema Validation", "Pydantic (structured LLM outputs)", "Free"],
            ["QA Framework", "Pytest + GitHub Actions CI", "Free"],
            ["Version Control", "Git + GitHub", "Free"],
        ],
    )
    add_para(
        doc,
        "11 of 12 tech-stack components are completely free. The ₹2,500 ceiling is allocated entirely "
        "to LLM inference via OpenRouter, with Ollama as a zero-cost local fallback.",
        italic=True,
    )

    # ===== 9. Expected Deliverable / POC =====
    add_heading(doc, "9. Expected Deliverable / POC", level=2)
    add_para(doc, "At the Week 17 Final Review, Psiog Kendra will demonstrate:")
    for b in [
        "Live multi-agent system deployed on Azure with Coordinator + 4 Specialist Agents (Data, DevOps, CRM, Docs)",
        "Natural-language queries answered across all four domains with grounded, cited responses",
        "AgentField vector memory pre-indexed with internal documentation (runbooks, architecture docs)",
        "All 12 mandatory test queries passing — including 4 cross-domain queries that span 2+ specialists",
        "Routing accuracy: 100% on all 12 defined test queries",
        "Hallucination rate documented (target <10%) via a Judge Agent that verifies each claim against its citation",
        "Pydantic-validated structured outputs (RoutingDecision, AgentResponse) for machine-verifiable QA",
        "Full evidence pack on Moodle: architecture, agent code, routing logs, hallucination report, cost breakdown, demo recording",
    ]:
        add_bullet(doc, b)

    add_para(doc, "Twelve mandatory test queries:", bold=True)
    add_table(
        doc,
        ["#", "Query", "Expected Domain(s)"],
        [
            ["1", "Did yesterday's ETL pipeline run successfully?", "Data Platform"],
            ["2", "What was the error in the last failed Databricks job?", "Data Platform"],
            ["3", "Did the payments service deployment pass all quality gates?", "DevOps"],
            ["4", "What was the last deployment date for the auth service?", "DevOps"],
            ["5", "What is the deal status for Acme Corp?", "CRM"],
            ["6", "Who is the account owner for TechStart Ltd?", "CRM"],
            ["7", "What is the runbook for a schema mismatch error?", "Docs"],
            ["8", "What does the architecture doc say about the ingestion pipeline?", "Docs"],
            ["9", "Did last night's pipeline failure affect any CRM customer sync?", "Data Platform + CRM"],
            ["10", "The ingestion job failed — is there a fix in the runbooks?", "Data Platform + Docs"],
            ["11", "What's the status of the latest deployment and are there known issues?", "DevOps + Docs"],
            ["12", "Give me a full status update — pipeline, deployments, and any open incidents", "All 4 domains"],
        ],
    )

    # ===== 10. QA Strategy =====
    add_heading(doc, "10. QA Strategy (Mandatory)", level=2)
    add_para(
        doc,
        "QA is not an afterthought — it is built into the Psiog Kendra architecture from Week 1, "
        "embedded in every weekly deliverable, and automated via GitHub Actions CI.",
        italic=True,
    )

    add_para(doc, "10.1 Test Approach", bold=True)
    add_para(
        doc,
        "Psiog Kendra is verified across four layers: unit tests for each @app.skill function, "
        "integration tests for each specialist end-to-end, a routing accuracy suite for the Coordinator, "
        "and full E2E tests against the deployed system.",
    )

    add_para(doc, "10.2 Test Types & Coverage", bold=True)
    add_table(
        doc,
        ["Layer", "What We Test", "How", "Target"],
        [
            ["Unit", "Each @app.skill() function in isolation", "Pytest with mocked API responses", "≥80% code coverage"],
            ["Integration", "Each specialist agent end-to-end (query → cited response)", "Pytest hitting live AgentField control plane", "All 4 agents covered"],
            ["Routing", "Coordinator routes each of the 12 test queries correctly", "Automated routing-accuracy suite (actual vs expected)", "100% on 12 queries"],
            ["E2E", "Full query lifecycle from user input to final cited response", "curl / test harness against deployed system", "12 of 12 passing"],
            ["Data Validation", "All doc chunks have valid metadata; no orphan chunks; relevant retrieval", "Indexing audit + ≥5 representative similarity queries", "Zero invalid chunks"],
            ["Regression", "Full suite on every push", "GitHub Actions CI", "100% pass on merge"],
        ],
    )

    add_para(doc, "10.3 AI-Assisted QA — Judge Agent for Hallucination Detection", bold=True)
    add_para(
        doc,
        "Every specialist returns a citations list (source system or document). A separate Judge Agent "
        "(LLM reasoner) takes the agent's answer and citations, fetches the raw cited source, verifies "
        "each claim against it, and flags any claim that cannot be grounded.",
    )
    p = doc.add_paragraph()
    r = p.add_run("Hallucination Rate = (Ungrounded claims / Total claims) × 100   ·   measured across all 12 test queries")
    r.font.name = "Consolas"
    r.font.size = Pt(10)

    add_para(doc, "The QA Reasoner additionally:", bold=True)
    for b in [
        "Evaluates routing decisions (LLM-as-judge) — given the query and the chosen domains, is the routing correct?",
        "Detects hallucinations — does the answer contain claims not supported by citations?",
        "Scores response quality on a 1–5 scale — completeness, accuracy, citation clarity.",
    ]:
        add_bullet(doc, b)

    add_para(doc, "10.4 QA Deliverables", bold=True)
    for b in [
        "Pytest suite with ≥80% coverage across routing, retrieval, integration, and response generation",
        "GitHub Actions CI running all tests on every push",
        "Automated QA Report per run: routing accuracy %, hallucination rate %, response quality scores",
        "Routing accuracy log: actual vs expected domain for each of the 12 test queries",
        "Hallucination report: claim-by-claim grounding verification across all 12 queries",
        "All QA evidence uploaded to Moodle at Week 10 (mid-term) and Week 17 (final)",
    ]:
        add_bullet(doc, b)

    # ===== 11. Risk & Mitigations =====
    add_heading(doc, "11. Risk & Mitigations", level=2)
    add_table(
        doc,
        ["#", "Risk", "Impact", "Mitigation", "Fallback"],
        [
            ["1", "OpenRouter spend approaches ₹2,500 before Week 17", "Could exceed fixed budget", "Daily token usage tracking; cap per-query token budget; cache LLM responses for repeat test runs", "Switch primary LLM to Ollama + Llama 4 Scout (local, ₹0) via env-var change"],
            ["2", "External LLM exposure of sensitive data in production", "Privacy / compliance risk", "Use only mock/synthetic data in capstone demo; document three production mitigations (Ollama-local, Azure OpenAI Private Endpoint, Data Minimisation)", "For production: layered approach — Data Minimisation + Azure OpenAI Private Link + Ollama high-sensitivity fallback"],
            ["3", "Hallucination rate above 10% target", "Reduces trust / fails QA target", "Force every specialist to return citations; Judge Agent verifies grounding; tighten system prompts to ground-or-refuse", "Constrain answer template to citation-bound fields only; refuse if no citation found"],
            ["4", "Routing accuracy below 100% on 12 queries", "Coordinator sends queries to wrong specialist", "Few-shot routing examples in system prompt; Pydantic-validated RoutingDecision; routing-accuracy CI gate", "Add a confirmation step for ambiguous queries; fall back to broadcast-and-merge"],
            ["5", "Databricks / HubSpot free-tier API limits", "Specialist agents may fail in demo", "Cache results during demo; back off on rate limits; pre-warm common queries", "Replace with realistic mock APIs returning recorded fixtures"],
            ["6", "Azure B1s VM free-tier hours exhausted", "Hosting could incur cost", "CloudWatch-style usage check; stop VM outside dev hours", "Move to local Docker Compose for the demo; redeploy fresh free-tier VM"],
            ["7", "AgentField vector retrieval returns irrelevant chunks", "Docs Agent gives incorrect citations", "Tune chunk size; metadata-filtered search; evaluate top_k empirically across 5 sample queries", "Add lightweight BM25 re-ranker over top_k results"],
        ],
    )

    # ===== 12. Semester Alignment =====
    add_heading(doc, "12. Semester Alignment", level=2)
    add_para(
        doc,
        "Psiog Kendra aligns with Semester 4 (Capstone — Multi-Agent GenAI, QA Mandatory, AI Core) because:",
    )
    for b in [
        "It is an end-to-end multi-agent system — Coordinator + four Specialist Agents — not a single-LLM wrapper.",
        "AI is the core engine, not a bolt-on: LLM-based intent routing, LLM-driven response generation per agent, embeddings for semantic search, and an LLM Judge for QA. Without these the system does not function.",
        "Every response is grounded — citations from live system data or indexed internal docs — directly addressing the hallucination concern that motivates the RFP.",
        "QA is mandatory and embedded: 12 mandatory test queries, automated routing-accuracy and hallucination measurements, ≥80% coverage, Judge Agent for AI-assisted QA.",
        "Budget is efficiently managed at ₹2,500 of the ₹2,500 ceiling — every non-LLM component runs on free tiers or open-source, and Ollama provides a complete ₹0 fallback.",
        "Real business impact: collapses 4–5 separate tools and tribal-knowledge lookups into a single natural-language interface that a Day-1 joiner can use.",
    ]:
        add_bullet(doc, b)

    # ===== Closing summary table =====
    add_heading(doc, "Evaluation Summary (Self-Mapped to RFP Criteria)", level=2)
    add_table(
        doc,
        ["Criterion", "Psiog Kendra's Approach", "Weight"],
        [
            ["Problem Understanding", "Addresses fragmentation, tribal knowledge, and hallucination risk — not a paraphrase of the RFP", "15%"],
            ["Solution Quality", "5-agent architecture (1 Coordinator + 4 Specialists), clear routing flow, structured cited responses", "20%"],
            ["AI Integration Depth", "LLM for routing, LLM per agent, embeddings for semantic search, LLM Judge for QA — AI at every layer", "20%"],
            ["QA Strategy", "12 test queries, routing accuracy, hallucination rate, Judge Agent, ≥80% coverage", "15%"],
            ["Budget & Cost Justification", "₹2,500 of ₹2,500 — all non-LLM on free tiers; Ollama ₹0 fallback retained", "15%"],
            ["Timeline Feasibility", "Week-by-week milestones; Week 10 + Week 17 checkpoints with concrete deliverables", "15%"],
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
